import gspread
from oauth2client.service_account import ServiceAccountCredentials
from tkinter import *
from tkinter import ttk, filedialog, messagebox
import os
from docx import Document
from docx.shared import Pt
import webbrowser
from docx2pdf import convert

def connection_to_bd():
    # Подключение к Google-таблице
    json_keyfile = 'credentials/credd.json'
    scope = ['https://spreadsheets.google.com/feeds','https://www.googleapis.com/auth/drive', 'https://www.googleapis.com/auth/spreadsheets']
    credentials = ServiceAccountCredentials.from_json_keyfile_name(json_keyfile, scope)
    client = gspread.authorize(credentials)

    main_spreadsheet_id = '' # ID таблицы
    groups_spreadsheet_id = ''
    supervisors_spreadsheet_id = ''
    students_spreadsheet_id = ''
    profile_spreadsheet_id = ''

    main_worksheet_name = 'наполнение шаблонов'  # Имя листа таблицы, где содержатся данные
    groups_worksheet_name = 'группы'  
    supervisors_worksheet_name = 'преподаватели'  
    students_worksheet_name = 'студенты'  
    profile_worksheet_name = 'профили'  
    direction_worksheet_name = 'направления'  
    
    main_worksheet = client.open_by_key(main_spreadsheet_id).worksheet(main_worksheet_name)
    groups_worksheet = client.open_by_key(groups_spreadsheet_id).worksheet(groups_worksheet_name)
    supervisors_worksheet = client.open_by_key(supervisors_spreadsheet_id).worksheet(supervisors_worksheet_name)
    students_worksheet = client.open_by_key(students_spreadsheet_id).worksheet(students_worksheet_name)
    profile_worksheet = client.open_by_key(profile_spreadsheet_id).worksheet(profile_worksheet_name)
    direction_worksheet = client.open_by_key(profile_spreadsheet_id).worksheet(direction_worksheet_name)
    
    
    status_txt = 'Доступ к базе данных получен'

    return main_worksheet, groups_worksheet,supervisors_worksheet, students_worksheet, profile_worksheet, direction_worksheet, status_txt, main_spreadsheet_id

def open_templ_folder():
    os.system(r'start templates')

def row_finder(data, searched_word, n):
    row_to_select = None
    for i, row in enumerate(data):
        if row[n] == searched_word:  
            row_to_select = i + 1  # Нумерация рядов в таблице начинается с 1
            break
    return row_to_select

def fio_format(fio):
    surname, name, patronymic = fio.split()
    name = name[0] + '.'
    patronymic = patronymic[0] + '. '
    fio_new = name + patronymic + surname
    return(fio_new)

def cell_finder(word, worksheet, col_num):
    data = worksheet.get_all_values()
    row = row_finder(data, word, 0)
    return worksheet.cell(row, col_num).value


def key_replace(fio, doc):
    s_fio = fio_format(fio)

    s_group = cell_finder(fio, main_worksheet, 2)

    supervisor = cell_finder(fio, main_worksheet, 4)
    s_supervisor = fio_format(supervisor)

    s_supervisor_post = cell_finder(supervisor, supervisors_worksheet, 2)

    s_theme = cell_finder(fio, main_worksheet, 3)

    s_profile = cell_finder(s_group, groups_worksheet, 2)

    s_direction = cell_finder(s_profile, profile_worksheet, 2)

    s_code = cell_finder(s_direction, direction_worksheet, 2)

    replace_dict = {
    '{fio}': s_fio,
    '{group}':s_group,
    '{supervisor}': s_supervisor,
    '{supervisor_post}': s_supervisor_post,
    '{theme}': s_theme,
    '{profile}': s_profile,
    '{direction}':s_direction,
    '{code}':s_code    
    # Добавьте нужные слова и их замены в словарь
    }
    # for paragraph in doc.paragraphs:
    #     for word, replacement in replace_dict.items():
    #         if word in paragraph.text:
    #             for run in paragraph.runs:
    #                 if word in run.text:
    #                     run.text = run.text.replace(word, replacement)
    #                     alignment = paragraph.alignment
    #                     paragraph.alignment = alignment
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for word, replacement in replace_dict.items():
                       if word in paragraph.text: 
                           for run in paragraph.runs:
                               if word in run.text:
                                    run.text = run.text.replace(word, replacement)
                                    alignment = cell.paragraphs[0].alignment
                                    cell.paragraphs[0].alignment = alignment

def groupmate_search(group):
    students = []
    data = main_worksheet.get_all_values()
    for i, row in enumerate(data):
        if row[1] == group: 
                row_to_select = i + 1
                student = main_worksheet.cell(row_to_select, 1).value
                students.append(student)
    return students

def save():
    if templates_box.get()!='' and lb.get(ANCHOR)!='':
        if var.get() == 0:
            file=f'{directory}/{templates_box.get()}'
            fio=lb.get(ANCHOR)
            doc = Document(file)
            key_replace(fio, doc)
            path = filedialog.asksaveasfilename(filetypes=[('Word Document', '.docx')])
            if path != '':
                    doc.save(str(path) + '.docx')
            saved(path)
        else:
            file=f'{directory}/{templates_box.get()}'
            group=lb.get(ANCHOR)
            students = groupmate_search(group)
            path = filedialog.asksaveasfilename(filetypes=[('Word Document', '.docx')])
            for i in students:
                doc = Document(file)
                key_replace(i, doc)
                if path != '':
                    direct, filename = os.path.split(path)
                    fio_changed = fio_format(i)
                    doc.save(str(direct)+ '/' + fio_changed + '.docx')
            if path != '':
                messagebox.showinfo("Сохранение", "Файлы успешно сохранёны(наверно)")
    elif templates_box.get()=='':
        messagebox.showinfo("Уведомление", "Вы не выбрали шаблон")
    else:
        messagebox.showinfo("Уведомление", "Вы не выбрали студента")

        

def saved(path):
    if path!='':
        full_path = str(path) + '.docx'
        if os.path.exists(full_path):
            messagebox.showinfo("Сохранение", "Файл успешно сохранён")

def checkkey(event): 
    if var.get() == 0:
        listt = fio_list
    else:
        listt = group_list
    value = event.widget.get() 
    if value == '': 
        data = listt
    else: 
        data = [] 
        for item in listt: 
            if value.lower() in item.lower(): 
                data.append(item)                 
    update(data) 
   
def update(data): 
    lb.delete(0, 'end') 
    for item in data: 
        lb.insert('end', item) 

def open_basee():
    webbrowser.open(f'https://docs.google.com/spreadsheets/d/{main_spreadsheet_id}')

def arr_fill(data, listt, n):
    for i, row in enumerate(data):
        if n == 0:
            if row[n]!='ФИО':
                listt.append(row[n])
        if n == 1:
            if row[n]!='Группа':
                listt.append(row[n])
    return list(set(listt))


def radio_b():
    if var.get() == 0:
        FIO.delete (0, 'end')
        txt_fio['text'] = 'ФИО студента'
        update(fio_list) 
        solo_b['border']='2'
        group_b['border']='0'
    else:
        FIO.delete (0, 'end')
        txt_fio['text'] = 'Номер группы'
        update(group_list) 
        solo_b['border']='0'
        group_b['border']='2'
    
    

fio_list=[]
group_list=[]

# Путь к директории с шаблонами DOCX
directory = './templates'

templates = [file for file in os.listdir(directory) if file.endswith('.docx')]

# создание GUI приложения
root = Tk()

root['bg'] = '#fafafa'
root.title('Генератор титулов')
root.geometry('600x450')

root.resizable(width=False, height=False)

txt_sh = Label(root, text='Выберите шаблон', font=20)
txt_sh.place(x=70,y=30)

templates_box = ttk.Combobox(values=templates,state='readonly', width=35)
templates_box.place(x=70,y=65)

txt_fio = Label(root, text='ФИО студента', font=20)
txt_fio.place(x=360,y=30)

FIO = Entry(root, width=28) 
FIO.place(x=360,y=65)
FIO.bind('<KeyRelease>', checkkey) 
  
lb = Listbox(root,height=8, width=28) 
lb.place(x=360,y=100)

btn = Button(root, text='Сохранить', width=25,height=3,command=save)
btn.place(x=360,y=320)

btn = Button(root, text='Открыть папку с шаблонами', width=25,height=3, command=open_templ_folder)
btn.place(x=70,y=320)

btn = Button(root, text='Открыть базу данных', width=25,height=3, command=open_basee)
btn.place(x=70,y=235)

frame = Frame(borderwidth=1, relief=SOLID, width=8, height=8, )
frame.place(x=390,y=426)

var = IntVar()
var.set(0)

solo_b = Radiobutton(text="Для одного студента", indicatoron=False, border=2, padx=3, pady=3,  width=18, height=2, variable=var,  value=0, command=radio_b)
group_b = Radiobutton(text="Для группы", variable=var, indicatoron=False, border=0, padx=3, pady=3,  width=18, height=2, value=1, command=radio_b )
solo_b.place(x=70,y=110)
group_b.place(x=70,y=160)

try:
    # Вызов функции для подключения
    main_worksheet, groups_worksheet,supervisors_worksheet, students_worksheet, profile_worksheet, direction_worksheet, status_txt, main_spreadsheet_id = connection_to_bd()
    main_ws_data = main_worksheet.get_all_values()
    # заполнение массивов для списка
    fio_list = arr_fill(main_ws_data, fio_list, 0)
    group_list = arr_fill(main_ws_data, group_list, 1)
    update(fio_list) 
except:
    status_txt='Ошибка подключения к базе данных'


if status_txt=='Доступ к базе данных получен':
    frame['bg']='green'
else:
    frame['bg']='red'

status = Label(root, text=status_txt, font="Arial 8")
status.place(x=400,y=420)

root.mainloop()
